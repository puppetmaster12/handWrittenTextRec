from docx import Document
from docx.shared import Inches

import datetime
import os

def test_savedoc():
    text = "This is a test document"
    t = datetime.datetime.now()
    current = t.strftime('%m-%d-%y')
    save_name = "./test_doc/test_doc_" + current + ".docx"
    document = Document()
    document.add_heading('Recognized text', 0)

    p = document.add_paragraph(text)
    document.add_page_break()
    document.save(save_name)

    assert os.path.exists(save_name) == 1
