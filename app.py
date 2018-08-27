from flask import Flask, render_template, request, json, jsonify, send_file
import tensorflow as tf
import numpy as np
import pandas as pd
import cv2
import os
import logging
import time
import datetime
from reportlab.pdfgen import canvas
from logging.handlers import RotatingFileHandler
from docx import Document
from docx.shared import Inches

from ocr.normalization import imageNorm, letterNorm
from ocr import page, words, charSeg
from ocr.helpers import implt, resize
from ocr.tfhelpers import Graph
from ocr.datahelpers import idx2char

current_milli_time = lambda: int(round(time.time() * 1000))

import Cycler

app = Flask(__name__)

if not app.debug:
    formatter = logging.Formatter(
        "[%(asctime)s] {%(pathname)s:%(lineno)d} %(levelname)s - %(message)s")
    handler = RotatingFileHandler('logs/errors.log', maxBytes=10000000, backupCount=5)
    handler.setLevel(logging.DEBUG)
    handler.setFormatter(formatter)
    app.logger.addHandler(handler)


MODEL_LOC = 'models/char-clas/en/CharClassifier'

@app.route("/")
def main():
    return render_template('index.html')

@app.route("/recognize", methods=['POST'])
def recognize():
    if request.method == 'POST':
        # start = current_milli_time()
        if 'image' not in request.files:
            print('No file part')
            return render_template('index.html')
        file = request.files['image']
        if file.filename == '':
            print('No selected file')
            return render_template('index.html')
        if file:
            file.save(os.path.join('images',file.filename))
            filename = file.filename
            # load the trained model
            charClass = Graph(MODEL_LOC)
            # load the saved image
            image = cv2.cvtColor(cv2.imread("images/"+filename), cv2.COLOR_BGR2RGB)
            crop = page.detection(image)
            bBoxes = words.detection(crop)
            cycler = Cycler.Cycler(crop,bBoxes,charClass)
            allWords = []
            for i in range(len(bBoxes)-1):
                allWords.append(cycler.idxImage(i))
                print(allWords[i])
    # stop = current_milli_time()
    # print(stop - start)

    return jsonify(allWords=allWords)

@app.route("/savePdf", methods=['POST'])
def savePdf():
    if request.method == 'POST':
        text = request.form['text']
        t = datetime.datetime.now()
        current = t.strftime('%m-%d-%y')
        save_name = "doc/pdf_" + current + ".pdf"
        c = canvas.Canvas(save_name)
        c.drawString(40, 800, text)
        c.save()
        return send_file(save_name, as_attachment=True)

@app.route("/saveDoc", methods=['POST'])
def saveDoc():
    if request.method == 'POST':
        text = request.form['textDoc']
        # print("text", text)
        t = datetime.datetime.now()
        current = t.strftime('%m-%d-%y')
        save_name = "doc/docx_" + current + ".docx"
        document = Document()
        document.add_heading('Recognized text', 0)

        p = document.add_paragraph(text)
        document.add_page_break()
        document.save(save_name)
        return send_file(save_name, as_attachment=True)

if __name__ == "__main__":
    app.jinja_env.auto_reload = True
    app.config['TEMPLATES_AUTO_RELOAD'] = True
    app.run(debug=True)
